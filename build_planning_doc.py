from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(__file__).parent / "AICE_BASIC_모의고사_플랫폼_기획안.docx"
NAVY = "183B56"; BLUE = "2878B5"; LIGHT = "EAF2F8"; PALE = "F5F7FA"; GRAY = "667085"; WHITE = "FFFFFF"; BLACK = "202124"; GREEN = "16794F"; GOLD = "A86B00"

def font(run, size=10.5, bold=False, color=BLACK, name="Malgun Gothic"):
    run.font.name = name; run.font.size = Pt(size); run.bold = bold; run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    return run

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for tag,val in [("top",top),("start",start),("bottom",bottom),("end",end)]:
        node=tcMar.find(qn(f"w:{tag}"))
        if node is None: node=OxmlElement(f"w:{tag}"); tcMar.append(node)
        node.set(qn("w:w"),str(val)); node.set(qn("w:type"),"dxa")

def set_cell_text(cell, text, bold=False, color=BLACK, size=9.2, align=None):
    cell.text=""; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.08
    if align is not None: p.alignment=align
    font(p.add_run(str(text)),size,bold,color); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)

def table(doc, headers, rows, widths=None, font_size=9.0):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i],h,True,WHITE,9.0,WD_ALIGN_PARAGRAPH.CENTER); shade(t.rows[0].cells[i],NAVY)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            set_cell_text(cells[i],v,False,BLACK,font_size,WD_ALIGN_PARAGRAPH.CENTER if i==0 else WD_ALIGN_PARAGRAPH.LEFT)
            if len(t.rows)%2==1: shade(cells[i],PALE)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t

def h(doc, text, level=1):
    p=doc.add_paragraph(style=f"Heading {level}"); p.add_run(text); return p

def para(doc, text="", bold_lead=None, color=BLACK, after=5):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.12
    if bold_lead and text.startswith(bold_lead):
        font(p.add_run(bold_lead),10.5,True,color); font(p.add_run(text[len(bold_lead):]),10.5,False,color)
    else: font(p.add_run(text),10.5,False,color)
    return p

def bullets(doc, items, numbered=False):
    style="List Number" if numbered else "List Bullet"
    for item in items:
        p=doc.add_paragraph(style=style); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.10
        font(p.add_run(item),10.2)

def callout(doc, label, text, fill=LIGHT, accent=BLUE):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; c=t.cell(0,0); shade(c,fill); margins(c,150,180,150,180)
    c.text=""; p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.15
    font(p.add_run(label+"  "),10.3,True,accent); font(p.add_run(text),10.3)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=Inches(.82); sec.bottom_margin=Inches(.75); sec.left_margin=Inches(.85); sec.right_margin=Inches(.85); sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)

styles=doc.styles
normal=styles["Normal"]; normal.font.name="Malgun Gothic"; normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string(BLACK)
normal._element.rPr.rFonts.set(qn("w:eastAsia"),"Malgun Gothic"); normal.paragraph_format.space_after=Pt(5); normal.paragraph_format.line_spacing=1.12
for name,size,color,before,after in [("Heading 1",16,BLUE,16,7),("Heading 2",13,NAVY,11,5),("Heading 3",11.5,NAVY,8,4)]:
    s=styles[name]; s.font.name="Malgun Gothic"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s._element.rPr.rFonts.set(qn("w:eastAsia"),"Malgun Gothic"); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(header.add_run("AICE BASIC 모의고사 플랫폼 | 서비스 기획안"),8.5,False,GRAY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); footer._p.append(fld)

# Cover
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(80); p.paragraph_format.space_after=Pt(10); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
font(p.add_run("SERVICE PLANNING DOCUMENT"),10,True,BLUE)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(8); font(p.add_run("AICE BASIC\n모의고사 플랫폼 기획안"),28,True,NAVY)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(28); font(p.add_run("AIDU 실습 연계 · 자동채점 · 영역별 학습진단"),14,False,GRAY)
callout(doc,"프로젝트 목표","실제 AICE BASIC 응시 흐름과 유사한 환경에서 AIDU 실습 결과를 답안으로 제출하고, 즉시 채점과 취약영역 분석을 제공하는 웹 기반 모의고사 서비스",LIGHT,BLUE)
for k,v in [("문서 버전","v1.0"),("작성일","2026년 7월 12일"),("기술 구성","Next.js · Supabase · Vercel"),("초기 콘텐츠","분류형 3세트 · 회귀형 3세트")]:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); font(p.add_run(f"{k}  "),10,True,NAVY); font(p.add_run(v),10,False,GRAY)
doc.add_page_break()

h(doc,"문서 개요",1)
para(doc,"본 문서는 AICE BASIC 대비 모의고사 플랫폼의 서비스 방향, 사용자 경험, 기능 요구사항, 데이터 구조, 채점 정책, 보안 원칙 및 단계별 구축 범위를 정의한다.")
h(doc,"목차",2)
table(doc,["구분","내용"],[("1","프로젝트 개요 및 시험 분석"),("2","서비스 목표와 사용자 정의"),("3","사용자 여정 및 화면 구성"),("4","기능 요구사항"),("5","채점 및 학습진단 정책"),("6","시스템·Supabase 설계"),("7","보안·개인정보·운영 정책"),("8","MVP 범위 및 개발 로드맵"),("9","검수 기준과 주요 의사결정")],[.65,5.65])
callout(doc,"핵심 원칙","사이트는 AIDU 자체를 대체하지 않는다. 사용자는 설치형 AIDU에서 데이터를 처리하고 모델을 실행하며, 웹 서비스는 시험 진행·답안 저장·자동채점·분석을 담당한다.","FFF7E6",GOLD)

h(doc,"1. 프로젝트 개요 및 시험 분석",1)
h(doc,"1.1 AICE BASIC 개요",2)
para(doc,"AICE는 KT가 개발하고 한국경제신문과 함께 주관하는 AI 활용능력 시험이다. BASIC 등급은 AI 입문자와 비전공자가 노코딩 환경에서 비즈니스 데이터를 분석하고 AI 모델을 만들어 문제를 해결하는 역량을 평가한다.")
table(doc,["항목","공식 안내 기준"],[("대상","AI 입문자, 비전공자"),("방식","온라인 비대면, 100% 노코딩 실기"),("문항·시간","15문항, 60분"),("합격 기준","100점 만점 중 80점 이상"),("실습 도구","설치형 노코딩 AI 도구 AIDU"),("데이터 유형","Tabular Data"),("유효기간","영구")],[1.55,4.75])
h(doc,"1.2 출제영역",2)
table(doc,["평가영역","주요 내용","문항","배점"],[("탐색적 데이터 분석","기초통계, 시각화, 상관관계, 이상치","5","30"),("데이터 전처리","결측치·이상치 처리, 인코딩, 정규화","5","30"),("AI 모델링","머신러닝·딥러닝 모델 학습","2","16"),("모델 성능평가","평가지표, 영향도, 개선, 시뮬레이션","3","24")],[1.55,3.25,.65,.65])
para(doc,"출처: AICE BASIC 공식 안내 - https://aice.study/info/aice/basic",color=GRAY,after=8)
h(doc,"1.3 보유 샘플 분석",2)
para(doc,"작업 폴더에는 분류형 3세트와 회귀형 3세트가 있으며, 각 세트는 해설 PDF, AIDU 실습 안내서, 실습 CSV로 구성된다. 세트별 문항 수는 약 11~15문항이며 일부 문항은 앞 문항에서 가공한 데이터나 저장한 모델을 이어서 사용한다.")
table(doc,["유형","주제·데이터","주요 평가요소"],[("분류 3세트","미디어 고객 콘텐츠 취향 예측 / moviegenre.csv","분류 판단, 결측치, 시각화, Accuracy·Precision·Recall·F1, 영향도, 시뮬레이션"),("회귀 3세트","통신 이용요금 예측 / amount 계열 CSV","회귀 판단, 기초통계, 이상치, R²·MAE·MSE, 영향도, 예측, 성능 개선")],[1.05,2.0,3.25])
callout(doc,"콘텐츠 운영 판단","초기 6세트는 샘플 모의고사로 제공하고, 이후 15문항·100점 구성의 문제은행형 실전 모의고사를 확장한다. 선행 문항이 있는 세트는 문제 순서를 임의로 섞지 않는다.")

h(doc,"2. 서비스 목표와 사용자 정의",1)
h(doc,"2.1 서비스 목표",2)
bullets(doc,["AIDU 사용 경험과 실제 시험 대응력을 높인다.","시험 종료 즉시 점수와 모의 합격 여부를 제공한다.","단순 총점이 아니라 영역·세부역량별 취약점을 진단한다.","관리자가 문제·정답·허용오차·해설을 코드 수정 없이 운영할 수 있게 한다.","응시 이력과 성장 추이를 축적해 반복 학습을 유도한다."])
h(doc,"2.2 핵심 사용자",2)
table(doc,["사용자","필요","제공 가치"],[("AI 입문자·비전공자","AIDU 조작 및 시험 흐름 경험","설치 안내, 단계형 문제, 즉시 해설"),("AICE BASIC 수험생","실전 시간 관리와 합격 가능성 확인","60분 시험, 자동채점, 합격 판정"),("교육 운영자·관리자","문제와 성적의 효율적 운영","문제은행, 통계, 공개 제어")],[1.55,2.25,2.5])
h(doc,"2.3 성공지표",2)
bullets(doc,["가입 후 첫 시험 시작률", "시험 완주율 및 평균 응시시간", "재응시율과 2회차 점수 향상률", "영역별 분석지 조회율", "문항별 정답률 및 이의제기 발생률", "AIDU 설치 안내에서 시험 시작까지의 이탈률"])

h(doc,"3. 사용자 여정 및 화면 구성",1)
h(doc,"3.1 전체 사용자 여정",2)
table(doc,["단계","사용자 행동","시스템 처리"],[("1. 가입","이메일·비밀번호·생년월일 입력","이메일 인증, 프로필 및 동의 저장"),("2. 시험 선택","분류/회귀 모의고사 선택","응시 가능 여부와 이전 기록 표시"),("3. 준비","AIDU 확인, CSV·가이드 다운로드","시험 자산 접근 기록"),("4. 응시","AIDU 실행 결과를 답안에 입력","타이머, 자동저장, 응답 검증"),("5. 제출","최종 제출 또는 시간 종료","서버 측 채점 및 답안 잠금"),("6. 결과","총점·합격·오답 확인","영역 및 세부 태그별 집계"),("7. 진단","그래프와 코멘트 확인","취약영역 및 추천 복습 생성")],[.8,2.25,3.25])
h(doc,"3.2 회원가입·인증",2)
bullets(doc,["필수 입력: 이메일, 비밀번호, 비밀번호 확인, 생년월일", "필수 동의: 이용약관, 개인정보 수집·이용", "Supabase Auth 이메일·비밀번호 인증 및 이메일 확인", "비밀번호 최소 8자 정책, 찾기·재설정·인증메일 재발송", "비밀번호 원문은 별도 저장하지 않으며 Supabase Auth에 위임", "생년월일과 동의 이력은 별도 테이블에 최소한으로 저장"])
h(doc,"3.3 주요 화면",2)
table(doc,["화면","핵심 구성"],[("랜딩","시험 소개, 이용 흐름, AIDU 안내, 가입·로그인 CTA, 비공식 학습서비스 고지"),("대시보드","시험 목록, 최근·최고·평균 점수, 응시 이력, 취약영역, 이어하기"),("시험 안내","과제 배경, 제한시간, 배점, 파일 다운로드, AIDU 확인, 시작"),("시험 응시","타이머, 문제 내비게이션, 답안 입력, 자동저장, 검토 표시"),("결과","총점, 모의 합격/불합격, 정오표, 영역별 점수, 응시시간"),("분석지","막대·레이더 차트, 평균 비교, 세부역량 진단, 보강 코멘트, 오답 해설"),("관리자","시험·문제·정답·허용오차·파일·해설·공개상태·통계 관리")],[1.25,5.05])

h(doc,"4. 기능 요구사항",1)
h(doc,"4.1 시험 진행",2)
bullets(doc,["시험 시작 시 서버 기준 시작·종료시간이 포함된 응시 세션 생성", "60분 카운트다운 및 시간 종료 자동 제출", "답안 변경 시 자동저장, 저장 상태 표시, 새로고침 복구", "미응답·응답완료·검토표시를 문제 번호에 구분", "최종 제출 전 미응답 수와 제출 불가역성 안내", "제출 완료 후 답안 수정 금지", "선행 문항 의존성을 유지하는 고정 순서 옵션"])
h(doc,"4.2 답안 유형",2)
table(doc,["유형","입력 UI","검증"],[("객관식 단일선택","라디오 버튼","선택지 ID 기준"),("정수","숫자 입력","쉼표·공백 정규화"),("소수","숫자 입력","지정 자릿수 또는 허용오차"),("백분율","숫자+단위 안내","% 제거 후 수치 비교"),("단위 포함","값/단위 분리 또는 형식 안내","k 등 단위 환산"),("문자·범주","선택형 우선","대소문자·공백 정규화")],[1.15,2.05,3.1])
h(doc,"4.3 결과와 재학습",2)
bullets(doc,["제출 직후 총점, 정답 수, 오답 수, 응시시간 표시", "80점 이상을 모의 합격으로 판정", "영역별 취득점수·만점·정답률 제공", "문항별 사용자 답안·정답·해설 제공", "틀린 문항만 다시 보기 및 전체 재응시", "최근 응시 대비 증감과 누적 추이 제공"])

h(doc,"5. 채점 및 학습진단 정책",1)
h(doc,"5.1 서버 측 채점 원칙",2)
para(doc,"정답은 브라우저로 내려보내지 않고 제출 시 서버 전용 로직에서 채점한다. 선택지 순서를 변경할 수 있으므로 객관식은 화면 번호가 아닌 선택지의 고유 ID를 기준으로 비교한다.")
table(doc,["채점방식","적용 예","처리"],[("정확 일치","객관식, 정수, 범주","정규화 후 완전 일치"),("반올림 일치","0.9536 등 성능지표","지정 소수 자릿수 반올림"),("절대 오차","AIDU 결과의 미세 차이","|입력-정답| ≤ 허용값"),("상대 오차","큰 예측금액","정답 대비 오차율 기준"),("복수 정답","버전별 인정값","등록된 인정 답 중 하나")],[1.25,2.0,3.05])
callout(doc,"필수 설계","AIDU 버전과 학습 환경에 따라 모델 성능값이 미세하게 달라질 수 있으므로 문항별 grading_type, decimal_places, tolerance를 설정한다.","FFF7E6",GOLD)
h(doc,"5.2 평가영역과 세부 태그",2)
para(doc,"각 문항에는 대분류 영역뿐 아니라 세부역량 태그를 부여한다. 이를 통해 '전처리가 약함'보다 '결측치 처리는 안정적이나 IQR·이상치 해석이 부족함'처럼 구체적인 진단을 제공한다.")
table(doc,["대분류","세부 태그 예시"],[("탐색적 데이터 분석","기초통계, 데이터 유형, 상관관계, 히트맵, 분포차트, 박스차트, 이상치"),("데이터 전처리","결측치, 대체전략, 인코딩, 스케일링, 변수 선택"),("AI 모델링","회귀·분류 판단, 머신러닝, 딥러닝, 파라미터 설정"),("모델 성능평가","Accuracy, Precision, Recall, F1, R², MAE, MSE, 영향도, 시뮬레이션, 개선")],[1.65,4.65])
h(doc,"5.3 진단 코멘트",2)
table(doc,["정답률","진단 수준","코멘트 방향"],[("80% 이상","안정","실전 시간 관리와 고난도 문항 권장"),("60~79%","보완","복합 조건과 결과 해석 연습"),("40~59%","취약","AIDU 메뉴 선택과 개념을 함께 복습"),("40% 미만","기초 필요","기초 개념 학습 후 쉬운 문항부터 재도전")],[1.15,1.15,4.0])
para(doc,"초기 버전은 일관성과 운영비를 고려해 규칙 기반 코멘트를 사용한다. 충분한 응시 데이터가 쌓인 뒤 생성형 AI 코칭을 선택적으로 검토한다.")
h(doc,"5.4 분석지 그래프",2)
bullets(doc,["영역별 취득점수/만점 막대그래프", "4개 평가영역 균형을 보여주는 레이더 차트", "본인 최근 응시 및 전체 평균 비교", "세부역량별 정답률과 취약도", "회차별 총점·영역점수 추이"])

h(doc,"6. 시스템 및 Supabase 설계",1)
h(doc,"6.1 기술 구성",2)
table(doc,["구성","권장 기술","역할"],[("웹","Next.js, TypeScript","페이지, 서버 라우트, SSR"),("UI","Tailwind CSS, shadcn/ui","반응형 시험·관리 화면"),("차트","Recharts","점수·추이 시각화"),("인증","Supabase Auth","이메일·비밀번호, 이메일 인증"),("DB","Supabase PostgreSQL","문제, 응시, 성적, 프로필"),("파일","Supabase Storage","CSV, 이미지, 가이드"),("배포","Vercel","자동 배포, 환경변수, 도메인")],[1.0,2.15,3.15])
h(doc,"6.2 핵심 데이터 모델",2)
table(doc,["테이블","주요 역할"],[("profiles","사용자 생년월일, 상태, 역할"),("user_consents","약관·개인정보 동의 버전과 시각"),("exams","시험명, 제한시간, 합격점, 상태"),("exam_sections","시험별 영역 및 만점"),("questions","문제, 유형, 배점, 영역, 순서, 선행관계"),("question_choices","객관식 보기"),("answer_keys","정답, 채점방식, 자릿수, 허용오차"),("exam_assets","CSV·가이드·문제 이미지"),("attempts","응시 시작·종료·상태·총점"),("attempt_answers","사용자 답, 채점 결과, 획득점수"),("section_results","영역별 집계 결과"),("diagnostic_rules","영역·태그·점수구간별 코멘트")],[1.75,4.55])
h(doc,"6.3 저장소와 접근",2)
bullets(doc,["exam-datasets: 응시용 CSV", "exam-assets: 문제 이미지 및 부속자료", "exam-guides: 공개 가능한 실습 가이드", "정답·내부 해설 원본은 공개 버킷에 저장하지 않음", "사용자별 응시답안은 RLS로 본인만 조회", "관리 기능은 별도 admin 역할과 서버 검증을 모두 적용"])
h(doc,"6.4 처리 흐름",2)
bullets(doc,["클라이언트가 답안을 입력하면 본인 응시 세션에 upsert", "최종 제출 요청에서 서버가 마감시간과 상태를 재검증", "서버가 정답키를 조회해 문항별 점수를 계산", "영역·태그별 결과와 진단을 생성해 트랜잭션으로 저장", "결과 페이지는 저장된 채점 결과만 조회"] , numbered=True)

h(doc,"7. 보안·개인정보·운영 정책",1)
h(doc,"7.1 개인정보",2)
bullets(doc,["수집 항목과 목적, 보유기간, 삭제 방법을 가입 화면과 처리방침에 명시", "생년월일은 필요한 목적 범위에서만 사용하고 관리자 화면에서도 최소 노출", "탈퇴 시 법적 보존 필요 항목을 제외한 프로필과 학습 데이터를 삭제 또는 익명화", "이메일 인증 여부와 동의 이력을 별도 기록", "운영 로그에 이메일·생년월일·답안 원문을 불필요하게 남기지 않음"])
h(doc,"7.2 시험·정답 보안",2)
bullets(doc,["answer_keys에 클라이언트 직접 SELECT 권한을 부여하지 않음", "Service Role 키는 Vercel 서버 환경에서만 사용", "공개 전 시험과 문제는 일반 사용자에게 차단", "제출 API에 중복 제출·변조·마감시간 검증", "관리자 변경 이력과 문제 공개 이력을 기록"])
h(doc,"7.3 저작권 및 표시",2)
bullets(doc,["AICE 명칭·로고·샘플 문제·해설·실습 가이드의 사용 권한 확인", "서비스 하단과 결과에 공식 주관 사이트가 아닌 학습용 모의고사임을 고지", "'합격'은 반드시 '모의고사 기준 합격'으로 표기", "공식 자료 원문 공개보다 허용된 범위 내 재구성 문제 사용을 우선"])

h(doc,"8. MVP 범위 및 개발 로드맵",1)
h(doc,"8.1 MVP 포함 범위",2)
table(doc,["영역","포함 기능"],[("계정","이메일·비밀번호 가입, 이메일 인증, 로그인, 재설정, 탈퇴"),("콘텐츠","분류 3세트, 회귀 3세트, CSV·안내자료"),("시험","60분 타이머, 자동저장, 이어하기, 자동제출"),("채점","객관식·숫자·소수·단위형, 허용오차, 서버 채점"),("결과","총점, 모의 합격, 정오표, 영역별 점수"),("분석","막대·레이더 차트, 규칙 기반 코멘트, 오답 해설"),("관리","시험·문제·정답·파일·공개상태·통계"),("배포","Vercel·Supabase 운영환경")],[1.15,5.15])
h(doc,"8.2 후속 확장",2)
bullets(doc,["문제은행 기반 랜덤 15문항 실전 시험", "응시자 평균·백분위와 난이도 보정", "회차별 성장 그래프와 맞춤 복습 세트", "교육기관·단체 계정 및 시험 코드", "관리자 통계 CSV 다운로드", "선택형 AI 상세 코칭"])
h(doc,"8.3 권장 구축 순서",2)
table(doc,["단계","작업","완료 기준"],[("1","6개 해설 PDF 문항 구조화·오류 검수","문항·정답·영역·배점·채점규칙 확정"),("2","정보구조·와이어프레임·디자인","핵심 화면 승인"),("3","Supabase 스키마·인증·RLS","계정과 권한 테스트 통과"),("4","시험·자동저장·타이머","새로고침·마감·복구 테스트"),("5","채점 엔진","전체 샘플 정답 회귀테스트"),("6","결과·분석지","영역 합계와 그래프 검증"),("7","관리자 기능","문제 추가·수정·공개 가능"),("8","Vercel 배포·운영점검","보안·성능·모바일 QA 통과")],[.55,2.9,2.85])

h(doc,"9. 검수 기준과 주요 의사결정",1)
h(doc,"9.1 출시 검수 기준",2)
bullets(doc,["6개 시험의 문제·보기·정답·해설이 원문과 대조 완료", "배점 합계와 영역별 합계가 100% 일치", "허용오차 경계값 채점 테스트 통과", "답안 자동저장, 새로고침 복구, 시간 종료 자동 제출 검증", "일반 사용자가 정답키와 타인 응시 결과에 접근할 수 없음", "모바일·태블릿·PC에서 시험 화면 사용 가능", "분석지의 총점·영역점수·그래프·코멘트가 동일 데이터로 계산", "개인정보 처리방침·약관·비공식 서비스 고지 적용"])
h(doc,"9.2 구현 전 확정할 의사결정",2)
table(doc,["항목","권장안","확정 필요 내용"],[("콘텐츠 사용권","공개 범위를 확인한 뒤 재구성","샘플 PDF·가이드·로고 사용 가능 범위"),("재응시 정책","횟수 제한 없이 이력 보존","무료/유료 및 회차 제한"),("해설 공개","제출 후 문항별 제공","정답만/상세 해설 범위"),("점수 비교","초기에는 본인 추이 중심","전체 평균·백분위 공개 시점"),("생년월일","연령 통계 목적 최소 수집","정확한 이용 목적·보유기간"),("관리 권한","admin 역할 기반","관리자 수와 승인 절차")],[1.2,2.1,3.0])
callout(doc,"다음 산출물","기획 승인 후 ① 샘플 6세트 문항 데이터 명세서, ② 화면 와이어프레임, ③ Supabase ERD·RLS 정책서, ④ 개발 백로그 순으로 구체화한다.",LIGHT,BLUE)

h(doc,"부록 A. 용어 정의",1)
table(doc,["용어","정의"],[("AIDU","AICE Junior·Basic 교육과 시험에 활용되는 설치형 노코딩 AI 실습 도구"),("응시 세션","시험 시작부터 제출까지의 사용자별 상태 단위"),("채점 규칙","정확 일치·반올림·허용오차 등 정답 비교 방식"),("평가영역","공식 출제범위에 대응하는 대분류"),("세부역량 태그","문항이 평가하는 구체 기술·개념"),("모의 합격","본 플랫폼의 기준점수에 따른 학습용 판정으로 실제 자격 취득과 무관")],[1.55,4.75])
h(doc,"부록 B. 참고자료",1)
bullets(doc,["AICE BASIC 공식 안내: https://aice.study/info/aice/basic", "유사 서비스 참고: https://ukbang.github.io/AICE_Basic_sample/", "작업 폴더 내 AICE BASIC 분류형·회귀형 해설 PDF 및 실습 CSV"])

# Keep headings with following paragraph and avoid orphaning table rows where practical
for p in doc.paragraphs:
    for r in p.runs:
        if not r.font.name: font(r,10.5)

doc.core_properties.title="AICE BASIC 모의고사 플랫폼 기획안"
doc.core_properties.subject="AIDU 실습 연계 자동채점 및 영역별 학습진단 서비스"
doc.core_properties.author=""
doc.save(OUT)
print(OUT)
